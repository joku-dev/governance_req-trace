import os
import tempfile
import unittest
from pathlib import Path

from openpyxl import load_workbook
from docx import Document


REPO_ROOT = Path(__file__).resolve().parent.parent
SRC_DIR = REPO_ROOT / "src"
if str(SRC_DIR) not in os.sys.path:
    os.sys.path.insert(0, str(SRC_DIR))

import devsecops_requirements_extractor as extractor


class ClassificationAndMetadataTests(unittest.TestCase):
    def test_policy_classification(self):
        sentence = "The organization policy shall be approved annually by the governance board."
        label, confidence = extractor.classify_requirement(sentence, "SHALL", "Governance")
        self.assertEqual(label, "Policy")
        self.assertIn(confidence, {"High", "Medium"})

    def test_control_classification(self):
        sentence = "The pipeline shall log and monitor build integrity events."
        label, confidence = extractor.classify_requirement(sentence, "SHALL", "CI/CD Controls")
        self.assertEqual(label, "Control")
        self.assertIn(confidence, {"High", "Medium"})

    def test_guidance_classification(self):
        sentence = "Teams should consider threat modeling for major architecture changes."
        label, confidence = extractor.classify_requirement(sentence, "SHOULD", "Secure Design")
        self.assertEqual(label, "Guidance")
        self.assertIn(confidence, {"High", "Medium"})

    def test_owner_and_evidence_inference(self):
        control_family = extractor.infer_control_family(
            "All privileged access must enforce MFA and RBAC.",
            "Access Control",
            "IAM, Access & Environment Security",
        )
        self.assertEqual(control_family, "IAM-ACCESS")

        owner = extractor.infer_owner_role(
            "The CISO must approve access exceptions.",
            "Governance",
            control_family,
        )
        self.assertEqual(owner, "CISO")

        evidence = extractor.infer_evidence_mapping(
            "Privileged access reviews shall be logged and retained.",
            "Control",
            control_family,
        )
        self.assertIn("Access review record", evidence)


class IntegrationTests(unittest.TestCase):
    def test_extract_requirements_from_docx_and_workbook_headers(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            input_doc = tmp_path / "security_controls.docx"
            output_xlsx = tmp_path / "out.xlsx"

            doc = Document()
            doc.add_heading("Security Policy", level=1)
            doc.add_paragraph(
                "The organization policy shall be reviewed annually. "
                "The pipeline shall log security events. "
                "Teams should perform threat modeling."
            )
            doc.save(input_doc)

            reqs, sources, docs, crossrefs = extractor.extract_requirements_from_documents([str(input_doc)])
            self.assertGreaterEqual(len(reqs), 3)
            self.assertEqual(len(reqs), len(sources))
            self.assertEqual(len(docs), 1)

            for req in reqs:
                self.assertIn(req.requirement_class, {"Policy", "Control", "Guidance"})
                self.assertIn(req.class_confidence, {"High", "Medium", "Low"})
                self.assertTrue(req.control_family)
                self.assertTrue(req.owner_role)
                self.assertTrue(req.evidence_mapping)

            extractor.create_workbook(reqs, sources, docs, crossrefs, str(output_xlsx))
            self.assertTrue(output_xlsx.exists())

            wb = load_workbook(filename=str(output_xlsx))
            ws_req = wb["Requirements_Master"]
            ws_src = wb["Source_Excerpts"]

            req_headers = [ws_req.cell(row=1, column=i).value for i in range(1, ws_req.max_column + 1)]
            src_headers = [ws_src.cell(row=1, column=i).value for i in range(1, ws_src.max_column + 1)]

            for required_col in [
                "Requirement_Class",
                "Class_Confidence",
                "Control_Family",
                "Owner_Role",
                "Evidence_Mapping",
            ]:
                self.assertIn(required_col, req_headers)
                self.assertIn(required_col, src_headers)


if __name__ == "__main__":
    unittest.main()
