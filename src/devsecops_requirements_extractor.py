import argparse
import os
import re
import sys
import traceback
from dataclasses import dataclass
from collections import defaultdict
from typing import List, Dict, Tuple, Optional

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
except Exception:
    tk = None

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo

# Optional Word COM support (best for .doc/.docx on Windows)
try:
    import win32com.client  # type: ignore
    HAS_WIN32 = True
except Exception:
    HAS_WIN32 = False

# Optional python-docx support (cross-platform .docx parsing)
try:
    from docx import Document as DocxDocument  # type: ignore
    HAS_PYDOCX = True
except Exception:
    HAS_PYDOCX = False

MODAL_PATTERNS = [
    ("PROHIBITED", re.compile(r"\b(SHALL\s+NOT|MUST\s+NOT|PROHIBITED|FORBIDDEN)\b", re.IGNORECASE)),
    ("SHALL", re.compile(r"\bSHALL\b", re.IGNORECASE)),
    ("MUST", re.compile(r"\bMUST\b", re.IGNORECASE)),
    ("SHOULD", re.compile(r"\bSHOULD\b", re.IGNORECASE)),
    ("MAY", re.compile(r"\bMAY\b", re.IGNORECASE)),
]

TOPIC_RULES = [
    ("Traceability", ["traceab", "trace", "requirement", "testcase", "artifact lineage", "author"]),
    ("Source Control & Change Integrity", ["branch", "commit", "merge", "version control", "repository", "review", "protected branch"]),
    ("Secure Coding", ["secure coding", "coding", "vulnerab", "defect", "static analysis", "sast"]),
    ("Dependencies, SBOM & Provenance", ["dependency", "sbom", "provenance", "component", "package", "third-party"]),
    ("Build Automation & Reproducibility", ["build", "pipeline", "reproduc", "artifact", "compile"]),
    ("Testing & Verification", ["test", "verification", "validation", "coverage", "unit test", "integration test"]),
    ("IAM, Access & Environment Security", ["access", "rbac", "mfa", "identity", "credential", "privilege", "environment"]),
    ("Secrets & Key Management", ["secret", "key", "certificate", "token", "credential"]),
    ("Vulnerability Management", ["vulnerability", "cve", "patch", "remediation", "severity"]),
    ("Governance Roles & Accountability", ["responsible", "accountable", "owner", "board", "authority", "governance"]),
    ("Waivers, Deviations & Exceptions", ["waiver", "deviation", "exception", "dispensation"]),
    ("Logging, Monitoring & Auditability", ["log", "audit", "monitor", "evidence", "record"]),
    ("Release & Deployment Control", ["release", "deploy", "promotion", "production", "approval"]),
]

CONTROL_FAMILY_BY_TOPIC = {
    "Traceability": "GOV-TRACE",
    "Source Control & Change Integrity": "CM-CHANGE",
    "Secure Coding": "APPSEC-SDLC",
    "Dependencies, SBOM & Provenance": "SUPPLY-CHAIN",
    "Build Automation & Reproducibility": "BUILD-CI",
    "Testing & Verification": "QA-VERIFICATION",
    "IAM, Access & Environment Security": "IAM-ACCESS",
    "Secrets & Key Management": "CRYPTO-SECRETS",
    "Vulnerability Management": "VULN-MGMT",
    "Governance Roles & Accountability": "GOV-ROLES",
    "Waivers, Deviations & Exceptions": "RISK-ACCEPTANCE",
    "Logging, Monitoring & Auditability": "AUDIT-LOGGING",
    "Release & Deployment Control": "RELEASE-DEPLOY",
    "Unclassified": "UNMAPPED",
}

FAMILY_OWNER_DEFAULTS = {
    "GOV-TRACE": "Requirements Manager",
    "CM-CHANGE": "Configuration Manager",
    "APPSEC-SDLC": "Application Security Lead",
    "SUPPLY-CHAIN": "Software Supply Chain Lead",
    "BUILD-CI": "DevOps Platform Owner",
    "QA-VERIFICATION": "QA/Test Manager",
    "IAM-ACCESS": "IAM Lead",
    "CRYPTO-SECRETS": "Security Engineering / PKI Owner",
    "VULN-MGMT": "Vulnerability Management Lead",
    "GOV-ROLES": "Governance Board Secretariat",
    "RISK-ACCEPTANCE": "Risk Manager",
    "AUDIT-LOGGING": "Security Operations Lead",
    "RELEASE-DEPLOY": "Release Manager",
    "UNMAPPED": "Control Owner (TBD)",
}

OWNER_ROLE_RULES = [
    ("CISO", ["ciso", "chief information security officer"]),
    ("Governance Board", ["board", "governance committee", "steering committee"]),
    ("Security Operations Lead", ["soc", "security operations"]),
    ("Application Security Lead", ["appsec", "secure coding"]),
    ("DevOps Platform Owner", ["devops", "platform team", "pipeline owner"]),
    ("QA/Test Manager", ["quality assurance", "qa team", "test manager"]),
    ("IAM Lead", ["identity team", "access management", "rbac", "mfa"]),
    ("Configuration Manager", ["change control board", "ccb", "configuration management"]),
    ("Release Manager", ["release manager", "deployment manager"]),
    ("Risk Manager", ["risk owner", "risk committee", "waiver"]),
]

CLASSIFICATION_KEYWORDS = {
    "Policy": [
        "policy", "standard", "baseline", "mandatory", "prohibited", "forbidden", "must", "shall", "shall not",
        "must not", "is required", "are required", "without exception", "governance",
    ],
    "Control": [
        "implement", "enforce", "verify", "monitor", "scan", "review", "approve", "audit", "log", "retain",
        "encrypt", "sign", "authenticate", "authorize", "segregate", "rotate", "test", "measure", "detect",
        "alert", "remediate", "block", "track", "ticket", "attestation",
    ],
    "Guidance": [
        "should", "may", "recommended", "guidance", "consider", "where feasible", "best practice", "can be",
        "typically", "advisory", "example",
    ],
}

CONTROL_FAMILY_KEYWORDS = [
    ("IAM-ACCESS", ["rbac", "mfa", "identity", "access", "privilege", "authentication", "authorization"]),
    ("CRYPTO-SECRETS", ["secret", "key", "certificate", "token", "vault", "kms", "hsm"]),
    ("VULN-MGMT", ["vulnerability", "cve", "patch", "remediation", "severity", "triage", "scanner"]),
    ("APPSEC-SDLC", ["sast", "dast", "secure coding", "code review", "threat model", "owasp"]),
    ("SUPPLY-CHAIN", ["sbom", "provenance", "dependency", "third-party", "component", "artifact signature"]),
    ("BUILD-CI", ["build", "pipeline", "reproducible", "ci", "compile"]),
    ("AUDIT-LOGGING", ["log", "logging", "audit trail", "monitor", "siem", "retention"]),
    ("CM-CHANGE", ["change", "branch", "commit", "merge", "repository", "version control"]),
    ("QA-VERIFICATION", ["test", "verification", "validation", "coverage", "unit test", "integration test"]),
    ("RELEASE-DEPLOY", ["release", "deployment", "promotion", "production", "rollout"]),
    ("RISK-ACCEPTANCE", ["waiver", "exception", "deviation", "dispensation", "risk acceptance"]),
    ("GOV-ROLES", ["responsible", "accountable", "owner", "authority", "board"]),
    ("GOV-TRACE", ["traceability", "trace", "artifact lineage", "requirement id"]),
]

EVIDENCE_KEYWORDS = [
    (["audit", "log", "monitor", "siem", "retention"], ["Audit log export", "SIEM report", "Retention configuration"]),
    (["approve", "review", "change", "ticket"], ["Approval record", "Change ticket", "Reviewer sign-off"]),
    (["test", "verification", "validation", "coverage"], ["Test report", "Coverage report", "Verification checklist"]),
    (["vulnerability", "cve", "scan", "patch"], ["Vulnerability scan report", "Patch record", "Remediation ticket"]),
    (["sbom", "provenance", "dependency"], ["SBOM export", "Provenance attestation", "Dependency review report"]),
    (["secret", "key", "certificate", "token"], ["Secrets inventory", "Key rotation log", "Certificate lifecycle record"]),
    (["identity", "access", "rbac", "mfa", "privilege"], ["Access review record", "IAM policy export", "MFA enforcement evidence"]),
    (["release", "deploy", "production", "promotion"], ["Release checklist", "Deployment approval", "Rollback test evidence"]),
]

FAMILY_EVIDENCE_DEFAULTS = {
    "IAM-ACCESS": ["Access review record", "IAM policy export"],
    "CRYPTO-SECRETS": ["Secrets inventory", "Key rotation log"],
    "VULN-MGMT": ["Vulnerability scan report", "Remediation ticket"],
    "APPSEC-SDLC": ["Static analysis report", "Secure code review record"],
    "SUPPLY-CHAIN": ["SBOM export", "Provenance attestation"],
    "BUILD-CI": ["Build log", "Pipeline configuration export"],
    "AUDIT-LOGGING": ["Audit log export", "Monitoring dashboard snapshot"],
    "CM-CHANGE": ["Change ticket", "Pull request approval"],
    "QA-VERIFICATION": ["Test report", "Verification checklist"],
    "RELEASE-DEPLOY": ["Release checklist", "Deployment approval"],
    "RISK-ACCEPTANCE": ["Signed waiver", "Risk acceptance record"],
    "GOV-ROLES": ["RACI matrix extract", "Governance meeting minutes"],
    "GOV-TRACE": ["Traceability matrix", "Requirement lineage export"],
    "UNMAPPED": ["Control execution evidence"],
}

HEADER_FILL = PatternFill("solid", fgColor="1F4E78")
HEADER_FONT = Font(color="FFFFFF", bold=True)
THIN = Side(style="thin", color="D9D9D9")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
LINK_FONT = Font(color="0563C1", underline="single")
WRAP_ALIGN = Alignment(vertical="top", wrap_text=True)


@dataclass
class RequirementRecord:
    req_id: str
    doc_id: str
    document: str
    baseline_level: str
    requirement_class: str
    class_confidence: str
    control_family: str
    owner_role: str
    evidence_mapping: str
    modal: str
    xref_group_id: str
    xref_topic: str
    src_id: str
    section_1: str
    section_2: str
    section_path: str
    source_paragraph_no: int
    lead_text: str
    requirement_text: str
    source_sentence_original: str
    extraction_type: str
    source_doc_path: str


@dataclass
class SourceExcerptRecord:
    src_id: str
    req_id: str
    doc_id: str
    document: str
    source_paragraph_no: int
    section_1: str
    section_2: str
    section_path: str
    lead_text: str
    source_paragraph_original: str
    source_sentence_extracted: str
    requirement_text: str
    requirement_class: str
    class_confidence: str
    control_family: str
    owner_role: str
    evidence_mapping: str
    modal: str
    xref_group_id: str
    source_doc_path: str


@dataclass
class DocumentInfo:
    doc_id: str
    document: str
    filename: str
    full_path: str
    requirement_count: int
    notes: str
    first_req_id: Optional[str]


@dataclass
class CrossRefRecord:
    xref_group_id: str
    topic: str
    description: str
    requirement_ids: List[str]
    documents: List[str]
    first_req_id: str


class WordReader:
    def __init__(self):
        self.app = None

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        try:
            if self.app is not None:
                self.app.Quit(False)
        except Exception:
            pass

    def _ensure_com_app(self):
        if self.app is not None:
            return
        self.app = win32com.client.Dispatch("Word.Application")
        self.app.Visible = False
        self.app.DisplayAlerts = 0

    def read_document(self, path: str) -> List[Dict[str, str]]:
        ext = os.path.splitext(path)[1].lower()
        if ext == ".docx":
            if HAS_PYDOCX:
                return self._read_docx(path)
            if HAS_WIN32:
                self._ensure_com_app()
                return self._read_with_com(path)
            raise RuntimeError(
                "Cannot read .docx on this system. Install dependency 'python-docx'."
            )
        if ext == ".doc":
            if not HAS_WIN32:
                raise RuntimeError(
                    f"Cannot read legacy .doc file '{os.path.basename(path)}' on this platform. "
                    "Use Windows with Microsoft Word installed, or convert the file to .docx."
                )
            self._ensure_com_app()
            return self._read_with_com(path)
        raise RuntimeError(
            f"Unsupported input format '{ext}'. Supported formats are .docx (all platforms) and .doc (Windows + Word)."
        )

    def _read_docx(self, path: str) -> List[Dict[str, str]]:
        doc = DocxDocument(path)
        items = []
        current_h1 = ""
        current_h2 = ""
        paragraph_no = 0

        for para in doc.paragraphs:
            raw = (para.text or "").replace("\x07", " ").strip()
            if not raw:
                continue
            paragraph_no += 1
            try:
                style_name = str(para.style.name)
            except Exception:
                style_name = ""

            normalized = normalize_space(raw)
            if is_heading(style_name, normalized):
                level = detect_heading_level(style_name, normalized)
                if level == 1:
                    current_h1 = normalized
                    current_h2 = ""
                elif level == 2:
                    current_h2 = normalized
                continue

            items.append(
                {
                    "paragraph_no": paragraph_no,
                    "text": normalized,
                    "section_1": current_h1,
                    "section_2": current_h2,
                }
            )
        return items

    def _read_with_com(self, path: str) -> List[Dict[str, str]]:
        doc = self.app.Documents.Open(path, ReadOnly=True)
        items = []
        current_h1 = ""
        current_h2 = ""
        paragraph_no = 0
        try:
            for para in doc.Paragraphs:
                raw = para.Range.Text.replace("\r", " ").replace("\x07", " ").strip()
                if not raw:
                    continue
                paragraph_no += 1
                try:
                    style_name = str(para.Range.Style.NameLocal)
                except Exception:
                    style_name = ""

                normalized = normalize_space(raw)
                if is_heading(style_name, normalized):
                    level = detect_heading_level(style_name, normalized)
                    if level == 1:
                        current_h1 = normalized
                        current_h2 = ""
                    elif level == 2:
                        current_h2 = normalized
                    continue

                items.append(
                    {
                        "paragraph_no": paragraph_no,
                        "text": normalized,
                        "section_1": current_h1,
                        "section_2": current_h2,
                    }
                )
        finally:
            doc.Close(False)
        return items


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def is_heading(style_name: str, text: str) -> bool:
    s = (style_name or "").lower()
    if "heading" in s or "überschrift" in s:
        return True
    if len(text) < 120 and not text.endswith(".") and text == text.strip():
        if re.match(r"^[A-Z0-9][A-Za-z0-9 .,:()\-/]{2,}$", text):
            return True
    return False


def detect_heading_level(style_name: str, text: str) -> int:
    s = (style_name or "").lower()
    if "heading 1" in s or "überschrift 1" in s:
        return 1
    if "heading 2" in s or "überschrift 2" in s:
        return 2
    if re.match(r"^\d+(\.\d+)?\s+", text):
        return 2 if "." in text.split()[0] else 1
    return 2


def split_sentences(text: str) -> List[str]:
    parts = re.split(r"(?<=[.;:])\s+(?=[A-Z0-9(])", text)
    return [normalize_space(p) for p in parts if normalize_space(p)]


def detect_modal(sentence: str) -> Optional[str]:
    for label, pattern in MODAL_PATTERNS:
        if pattern.search(sentence):
            return label
    return None


def _contains_keyword(text: str, keyword: str) -> bool:
    escaped = re.escape(keyword)
    pattern = escaped
    if keyword and keyword[0].isalnum():
        pattern = r"\b" + pattern
    if keyword and keyword[-1].isalnum():
        pattern = pattern + r"\b"
    return re.search(pattern, text, re.IGNORECASE) is not None


def _has_any_keyword(text: str, keywords: List[str]) -> bool:
    return any(_contains_keyword(text, keyword) for keyword in keywords)


def _keyword_hits(text: str, keywords: List[str]) -> int:
    hits = 0
    for keyword in keywords:
        if _contains_keyword(text, keyword):
            hits += 1
    return hits


def _unique_join(values: List[str]) -> str:
    seen = set()
    ordered = []
    for value in values:
        if value not in seen:
            ordered.append(value)
            seen.add(value)
    return "; ".join(ordered)


def classify_requirement(sentence: str, modal: str, section_path: str) -> Tuple[str, str]:
    context = normalize_space(f"{section_path} {sentence}")
    scores = defaultdict(float)

    scores["Policy"] += 1.4 * _keyword_hits(context, CLASSIFICATION_KEYWORDS["Policy"])
    scores["Control"] += 1.6 * _keyword_hits(context, CLASSIFICATION_KEYWORDS["Control"])
    scores["Guidance"] += 1.4 * _keyword_hits(context, CLASSIFICATION_KEYWORDS["Guidance"])

    if modal in {"SHALL", "MUST", "PROHIBITED"}:
        scores["Policy"] += 2.0
        scores["Control"] += 1.2
    elif modal in {"SHOULD", "MAY"}:
        scores["Guidance"] += 2.0
        scores["Policy"] -= 0.5
        scores["Control"] -= 0.3

    if _has_any_keyword(context, ["verify", "monitor", "review", "approve", "scan", "test", "log", "retain"]):
        scores["Control"] += 1.8

    if _has_any_keyword(context, ["policy", "standard", "mandatory", "without exception"]):
        scores["Policy"] += 1.5

    ranked = sorted(
        [("Policy", scores["Policy"]), ("Control", scores["Control"]), ("Guidance", scores["Guidance"])],
        key=lambda x: x[1],
        reverse=True,
    )
    best_label, best_score = ranked[0]
    second_score = ranked[1][1]
    margin = best_score - second_score

    if margin >= 2.0:
        confidence = "High"
    elif margin >= 0.8:
        confidence = "Medium"
    else:
        confidence = "Low"

    return best_label, confidence


def infer_topic(sentence: str) -> Tuple[str, str]:
    for topic, keywords in TOPIC_RULES:
        if _has_any_keyword(sentence, keywords):
            xref_group_id = f"XREF-{TOPIC_RULES.index((topic, keywords)) + 1:03d}"
            return xref_group_id, topic
    return "XREF-999", "Unclassified"


def infer_control_family(sentence: str, section_path: str, xref_topic: str) -> str:
    topic_family = CONTROL_FAMILY_BY_TOPIC.get(xref_topic)
    if topic_family and topic_family != "UNMAPPED":
        return topic_family

    context = normalize_space(f"{section_path} {sentence}")
    for family, keywords in CONTROL_FAMILY_KEYWORDS:
        if _has_any_keyword(context, keywords):
            return family
    return "UNMAPPED"


def infer_owner_role(sentence: str, section_path: str, control_family: str) -> str:
    context = normalize_space(f"{section_path} {sentence}")
    for owner_role, keywords in OWNER_ROLE_RULES:
        if _has_any_keyword(context, keywords):
            return owner_role
    return FAMILY_OWNER_DEFAULTS.get(control_family, "Control Owner (TBD)")


def infer_evidence_mapping(sentence: str, requirement_class: str, control_family: str) -> str:
    context = sentence
    evidence_candidates: List[str] = []
    for keywords, evidence_items in EVIDENCE_KEYWORDS:
        if _has_any_keyword(context, keywords):
            evidence_candidates.extend(evidence_items)

    if not evidence_candidates:
        evidence_candidates.extend(FAMILY_EVIDENCE_DEFAULTS.get(control_family, ["Control execution evidence"]))

    if requirement_class == "Guidance":
        evidence_candidates.append("Adoption decision record")

    return _unique_join(evidence_candidates)


def infer_doc_id(filename: str, used: set) -> str:
    base = os.path.splitext(os.path.basename(filename))[0]
    tokens = re.findall(r"[A-Za-z]{2,}", base)
    preferred = None
    for token in tokens:
        up = token.upper()
        if len(up) in (3, 4):
            preferred = up
            break
    if not preferred:
        preferred = (re.sub(r"[^A-Za-z]", "", base).upper() or "DOC")[:3]
    doc_id = preferred[:4]
    n = 1
    original = doc_id
    while doc_id in used:
        n += 1
        doc_id = f"{original[:3]}{n}"
    used.add(doc_id)
    return doc_id


def infer_baseline_level(text: str) -> str:
    match = re.search(r"\bL([1-5])\b", text, re.IGNORECASE)
    if match:
        return f"L{match.group(1)}"
    return "N/A"


def document_title_from_filename(path: str) -> str:
    base = os.path.splitext(os.path.basename(path))[0]
    return normalize_space(base.replace("_", " ").replace("-", " "))


def extract_requirements_from_documents(paths: List[str]) -> Tuple[List[RequirementRecord], List[SourceExcerptRecord], List[DocumentInfo], List[CrossRefRecord]]:
    reqs: List[RequirementRecord] = []
    sources: List[SourceExcerptRecord] = []
    docs: List[DocumentInfo] = []
    cross_map: Dict[str, Dict[str, object]] = {}
    used_doc_ids = set()

    with WordReader() as reader:
        for path in paths:
            doc_id = infer_doc_id(path, used_doc_ids)
            document_name = document_title_from_filename(path)
            paragraphs = reader.read_document(path)
            local_count = 0
            first_req_id = None
            notes = "Auto-extracted from selected Word document. Review and curate results."

            for item in paragraphs:
                sentences = split_sentences(item["text"])
                for sent in sentences:
                    modal = detect_modal(sent)
                    if not modal:
                        continue
                    local_count += 1
                    req_id = f"{doc_id}-REQ-{local_count:03d}"
                    src_id = f"{doc_id}-SRC-{local_count:03d}"
                    if first_req_id is None:
                        first_req_id = req_id

                    xref_group_id, xref_topic = infer_topic(sent)
                    baseline_level = infer_baseline_level(sent + " " + item.get("section_1", ""))
                    section_1 = item.get("section_1", "")
                    section_2 = item.get("section_2", "")
                    section_path = " > ".join([x for x in [section_1, section_2] if x])
                    requirement_class, class_confidence = classify_requirement(sent, modal, section_path)
                    control_family = infer_control_family(sent, section_path, xref_topic)
                    owner_role = infer_owner_role(sent, section_path, control_family)
                    evidence_mapping = infer_evidence_mapping(sent, requirement_class, control_family)

                    reqs.append(
                        RequirementRecord(
                            req_id=req_id,
                            doc_id=doc_id,
                            document=document_name,
                            baseline_level=baseline_level,
                            requirement_class=requirement_class,
                            class_confidence=class_confidence,
                            control_family=control_family,
                            owner_role=owner_role,
                            evidence_mapping=evidence_mapping,
                            modal=modal,
                            xref_group_id=xref_group_id,
                            xref_topic=xref_topic,
                            src_id=src_id,
                            section_1=section_1,
                            section_2=section_2,
                            section_path=section_path,
                            source_paragraph_no=item["paragraph_no"],
                            lead_text="",
                            requirement_text=sent,
                            source_sentence_original=sent,
                            extraction_type="direct" if modal in {"SHALL", "MUST", "PROHIBITED"} else "heuristic",
                            source_doc_path=path,
                        )
                    )

                    sources.append(
                        SourceExcerptRecord(
                            src_id=src_id,
                            req_id=req_id,
                            doc_id=doc_id,
                            document=document_name,
                            source_paragraph_no=item["paragraph_no"],
                            section_1=section_1,
                            section_2=section_2,
                            section_path=section_path,
                            lead_text="",
                            source_paragraph_original=item["text"],
                            source_sentence_extracted=sent,
                            requirement_text=sent,
                            requirement_class=requirement_class,
                            class_confidence=class_confidence,
                            control_family=control_family,
                            owner_role=owner_role,
                            evidence_mapping=evidence_mapping,
                            modal=modal,
                            xref_group_id=xref_group_id,
                            source_doc_path=path,
                        )
                    )

                    if xref_group_id not in cross_map:
                        cross_map[xref_group_id] = {
                            "topic": xref_topic,
                            "description": f"Auto-grouped requirements for topic '{xref_topic}' with inferred control-family metadata.",
                            "requirement_ids": [],
                            "documents": set(),
                            "first_req_id": req_id,
                        }
                    cross_map[xref_group_id]["requirement_ids"].append(req_id)
                    cross_map[xref_group_id]["documents"].add(doc_id)

            docs.append(
                DocumentInfo(
                    doc_id=doc_id,
                    document=document_name,
                    filename=os.path.basename(path),
                    full_path=path,
                    requirement_count=local_count,
                    notes=notes,
                    first_req_id=first_req_id,
                )
            )

    crossrefs = []
    for xref_group_id in sorted(cross_map.keys()):
        payload = cross_map[xref_group_id]
        crossrefs.append(
            CrossRefRecord(
                xref_group_id=xref_group_id,
                topic=str(payload["topic"]),
                description=str(payload["description"]),
                requirement_ids=list(payload["requirement_ids"]),
                documents=sorted(list(payload["documents"])),
                first_req_id=str(payload["first_req_id"]),
            )
        )

    return reqs, sources, docs, crossrefs


def style_header(ws, row: int, start_col: int, end_col: int):
    for c in range(start_col, end_col + 1):
        cell = ws.cell(row=row, column=c)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.border = BORDER
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def style_body(ws, start_row: int, end_row: int, start_col: int, end_col: int):
    for r in range(start_row, end_row + 1):
        for c in range(start_col, end_col + 1):
            cell = ws.cell(r, c)
            cell.border = BORDER
            cell.alignment = WRAP_ALIGN


def add_table(ws, start_cell: str, end_cell: str, table_name: str):
    tab = Table(displayName=table_name, ref=f"{start_cell}:{end_cell}")
    style = TableStyleInfo(name="TableStyleMedium2", showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False)
    tab.tableStyleInfo = style
    ws.add_table(tab)


def set_column_widths(ws, widths: Dict[str, int]):
    for col, width in widths.items():
        ws.column_dimensions[col].width = width


def create_workbook(reqs: List[RequirementRecord], sources: List[SourceExcerptRecord], docs: List[DocumentInfo], crossrefs: List[CrossRefRecord], out_path: str):
    wb = Workbook()
    default = wb.active
    wb.remove(default)

    ws_readme = wb.create_sheet("README")
    ws_req = wb.create_sheet("Requirements_Master")
    ws_src = wb.create_sheet("Source_Excerpts")
    ws_xref = wb.create_sheet("Cross_Reference_Map")
    ws_docs = wb.create_sheet("Documents")
    class_counts = defaultdict(int)
    for req in reqs:
        class_counts[req.requirement_class] += 1

    # README
    readme_rows = [
        ["DevSecOps Requirements Extraction Workbook"],
        [],
        ["Purpose", "Consolidated extraction of requirements from selected Word documents."],
        ["How to use", "Use Requirements_Master as the primary filterable list. 'Source link' jumps to the extracted source excerpt. 'XREF link' jumps to the thematic cross-reference cluster. In Documents, the filename opens the original selected document."],
        ["Scope", "This script extracts requirement-like statements heuristically using modal verbs such as SHALL, MUST, SHOULD, MAY, SHALL NOT, and MUST NOT. It additionally classifies each row as Policy/Control/Guidance and enriches owner/control-family/evidence metadata."],
        ["Important", "Automatic extraction is never legally or semantically complete. Review each row before using it for governance, compliance, or audit."],
        [],
        ["Summary"],
        ["Total extracted rows", len(reqs)],
        ["Unique source documents", len(docs)],
        ["Rows with SHALL / MUST / etc.", len(reqs)],
        ["Rows classified as Policy", class_counts["Policy"]],
        ["Rows classified as Control", class_counts["Control"]],
        ["Rows classified as Guidance", class_counts["Guidance"]],
    ]
    for row in readme_rows:
        ws_readme.append(row)
    ws_readme["A1"].font = Font(size=14, bold=True)
    ws_readme["A1"].fill = HEADER_FILL
    ws_readme["A1"].font = Font(size=14, bold=True, color="FFFFFF")
    ws_readme["A1"].alignment = Alignment(horizontal="center")
    ws_readme.merge_cells("A1:D1")
    set_column_widths(ws_readme, {"A": 24, "B": 100, "C": 24, "D": 24})

    # Requirements_Master
    req_headers = [
        "REQ_ID", "Doc_ID", "Document", "Baseline_Level", "Requirement_Class", "Class_Confidence", "Control_Family",
        "Owner_Role", "Evidence_Mapping", "Modal", "XREF_Group_ID", "XREF_Topic", "SRC_ID", "Source_Link", "XREF_Link",
        "Section_1", "Section_2", "Section_Path", "Source_Paragraph_No", "Lead_Text", "Requirement_Text",
        "Source_Sentence_Original", "Extraction_Type"
    ]
    ws_req.append(req_headers)
    req_col = {name: idx + 1 for idx, name in enumerate(req_headers)}

    req_row_by_id = {}
    for idx, r in enumerate(reqs, start=2):
        ws_req.append([
            r.req_id, r.doc_id, r.document, r.baseline_level, r.requirement_class, r.class_confidence, r.control_family,
            r.owner_role, r.evidence_mapping, r.modal, r.xref_group_id, r.xref_topic, r.src_id, "Go to source",
            "Go to xref", r.section_1, r.section_2, r.section_path, r.source_paragraph_no, r.lead_text,
            r.requirement_text, r.source_sentence_original, r.extraction_type
        ])
        req_row_by_id[r.req_id] = idx

    # Source_Excerpts
    src_headers = [
        "SRC_ID", "REQ_ID", "Doc_ID", "Document", "Source_Paragraph_No", "Section_1", "Section_2", "Section_Path",
        "Requirement_Class", "Class_Confidence", "Control_Family", "Owner_Role", "Evidence_Mapping", "Lead_Text",
        "Source_Paragraph_Original", "Source_Sentence_Extracted", "Requirement_Text", "Modal", "XREF_Group_ID",
        "Requirement_Link", "Document_Link"
    ]
    ws_src.append(src_headers)
    src_col = {name: idx + 1 for idx, name in enumerate(src_headers)}
    src_row_by_id = {}
    for idx, s in enumerate(sources, start=2):
        ws_src.append([
            s.src_id, s.req_id, s.doc_id, s.document, s.source_paragraph_no, s.section_1, s.section_2, s.section_path,
            s.requirement_class, s.class_confidence, s.control_family, s.owner_role, s.evidence_mapping, s.lead_text,
            s.source_paragraph_original, s.source_sentence_extracted, s.requirement_text, s.modal, s.xref_group_id,
            "Back to requirement", "Open document"
        ])
        src_row_by_id[s.src_id] = idx

    # Cross reference
    xref_headers = ["XREF_Group_ID", "Topic", "Description", "Requirement_Count", "Requirement_IDs", "Documents", "First_REQ_ID", "First_REQ_Link"]
    ws_xref.append(xref_headers)
    xref_row_by_id = {}
    for idx, x in enumerate(crossrefs, start=2):
        ws_xref.append([
            x.xref_group_id, x.topic, x.description, len(x.requirement_ids), ", ".join(x.requirement_ids), ", ".join(x.documents), x.first_req_id, "Go to first requirement"
        ])
        xref_row_by_id[x.xref_group_id] = idx

    # Documents
    doc_headers = ["Doc_ID", "Document", "Filename", "Requirement_Count", "Purpose / Notes", "First_REQ_ID", "First_REQ_Link", "Full_Path"]
    ws_docs.append(doc_headers)
    for d in docs:
        ws_docs.append([d.doc_id, d.document, d.filename, d.requirement_count, d.notes, d.first_req_id or "", "Go to first requirement" if d.first_req_id else "", d.full_path])

    # Styling + tables
    for ws, headers in [
        (ws_req, req_headers), (ws_src, src_headers), (ws_xref, xref_headers), (ws_docs, doc_headers)
    ]:
        style_header(ws, 1, 1, len(headers))
        if ws.max_row >= 2:
            style_body(ws, 2, ws.max_row, 1, len(headers))
        ws.freeze_panes = "A2"
        add_table(ws, "A1", f"{get_column_letter(len(headers))}{ws.max_row}", f"T_{ws.title.replace('_', '')}")

    set_column_widths(ws_req, {
        "A": 14, "B": 10, "C": 34, "D": 12, "E": 16, "F": 14, "G": 16, "H": 28, "I": 40,
        "J": 10, "K": 14, "L": 28, "M": 14, "N": 14, "O": 14, "P": 24, "Q": 24, "R": 36,
        "S": 18, "T": 18, "U": 70, "V": 70, "W": 16,
    })
    set_column_widths(ws_src, {
        "A": 14, "B": 14, "C": 10, "D": 34, "E": 18, "F": 24, "G": 24, "H": 36, "I": 16,
        "J": 14, "K": 16, "L": 28, "M": 40, "N": 18, "O": 90, "P": 70, "Q": 70, "R": 12,
        "S": 14, "T": 18, "U": 18,
    })
    set_column_widths(ws_xref, {"A": 14, "B": 28, "C": 50, "D": 18, "E": 70, "F": 18, "G": 14, "H": 20})
    set_column_widths(ws_docs, {"A": 10, "B": 34, "C": 42, "D": 18, "E": 54, "F": 14, "G": 20, "H": 90})

    # Hyperlinks
    for row in range(2, ws_req.max_row + 1):
        src_id = ws_req.cell(row=row, column=req_col["SRC_ID"]).value
        xref_id = ws_req.cell(row=row, column=req_col["XREF_Group_ID"]).value
        if src_id in src_row_by_id:
            c = ws_req.cell(row=row, column=req_col["Source_Link"])
            c.hyperlink = f"#Source_Excerpts!A{src_row_by_id[src_id]}"
            c.font = LINK_FONT
        if xref_id in xref_row_by_id:
            c = ws_req.cell(row=row, column=req_col["XREF_Link"])
            c.hyperlink = f"#Cross_Reference_Map!A{xref_row_by_id[xref_id]}"
            c.font = LINK_FONT

    for row in range(2, ws_src.max_row + 1):
        req_id = ws_src.cell(row=row, column=src_col["REQ_ID"]).value
        if req_id in req_row_by_id:
            c = ws_src.cell(row=row, column=src_col["Requirement_Link"])
            c.hyperlink = f"#Requirements_Master!A{req_row_by_id[req_id]}"
            c.font = LINK_FONT
        doc_path = None
        src_id = ws_src.cell(row=row, column=src_col["SRC_ID"]).value
        src_obj = next((s for s in sources if s.src_id == src_id), None)
        if src_obj:
            doc_path = src_obj.source_doc_path
        if doc_path:
            c = ws_src.cell(row=row, column=src_col["Document_Link"])
            c.hyperlink = doc_path
            c.font = LINK_FONT

    for row in range(2, ws_xref.max_row + 1):
        req_id = ws_xref[f"G{row}"].value
        if req_id in req_row_by_id:
            c = ws_xref[f"H{row}"]
            c.hyperlink = f"#Requirements_Master!A{req_row_by_id[req_id]}"
            c.font = LINK_FONT

    for row in range(2, ws_docs.max_row + 1):
        filename_cell = ws_docs[f"C{row}"]
        full_path = ws_docs[f"H{row}"].value
        first_req_id = ws_docs[f"F{row}"].value
        if full_path:
            filename_cell.hyperlink = full_path
            filename_cell.font = LINK_FONT
        if first_req_id in req_row_by_id:
            c = ws_docs[f"G{row}"]
            c.hyperlink = f"#Requirements_Master!A{req_row_by_id[first_req_id]}"
            c.font = LINK_FONT

    wb.save(out_path)


def choose_files_gui() -> Tuple[List[str], str]:
    if tk is None:
        raise RuntimeError("tkinter is not available.")
    root = tk.Tk()
    root.withdraw()
    paths = filedialog.askopenfilenames(
        title="Select Word documents",
        filetypes=[("Word documents", "*.docx *.doc"), ("All files", "*.*")],
    )
    if not paths:
        return [], ""
    out_path = filedialog.asksaveasfilename(
        title="Save generated workbook as",
        defaultextension=".xlsx",
        initialfile="devsecops_requirements_extraction_generated.xlsx",
        filetypes=[("Excel Workbook", "*.xlsx")],
    )
    return list(paths), out_path


def parse_cli_args(argv: List[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Extract requirement-like statements from Word documents and generate an Excel workbook."
        )
    )
    parser.add_argument(
        "paths",
        nargs="*",
        help="Input Word files (.docx on all platforms, .doc only on Windows with Word).",
    )
    parser.add_argument(
        "-o",
        "--output",
        dest="output",
        help="Output .xlsx file path (default: ./devsecops_requirements_extraction_generated.xlsx).",
    )
    return parser.parse_args(argv[1:])


def main(argv: List[str]) -> int:
    try:
        args = parse_cli_args(argv)
        if args.paths:
            paths = [os.path.abspath(p) for p in args.paths if os.path.isfile(p)]
            if not paths:
                print("No valid input files provided.")
                return 2
            out_path = args.output or os.path.join(os.getcwd(), "devsecops_requirements_extraction_generated.xlsx")
        else:
            if tk is None:
                print(
                    "No input files were provided and tkinter GUI is unavailable.\n"
                    "Use CLI mode, e.g.:\n"
                    "  python src/devsecops_requirements_extractor.py input1.docx input2.docx -o output.xlsx"
                )
                return 2
            paths, out_path = choose_files_gui()
            if not paths or not out_path:
                print("Operation cancelled.")
                return 1

        reqs, sources, docs, crossrefs = extract_requirements_from_documents(paths)
        create_workbook(reqs, sources, docs, crossrefs, out_path)
        print(f"Workbook created: {out_path}")
        print(f"Documents processed: {len(docs)}")
        print(f"Requirements extracted: {len(reqs)}")
        if tk is not None:
            try:
                root = tk.Tk(); root.withdraw()
                messagebox.showinfo("Completed", f"Workbook created successfully:\n{out_path}\n\nRequirements extracted: {len(reqs)}")
                root.destroy()
            except Exception:
                pass
        return 0
    except Exception as exc:
        print("ERROR:", exc)
        traceback.print_exc()
        if tk is not None:
            try:
                root = tk.Tk(); root.withdraw()
                messagebox.showerror("Error", f"The program failed:\n{exc}")
                root.destroy()
            except Exception:
                pass
        return 99


if __name__ == "__main__":
    sys.exit(main(sys.argv))
